from docx import Document
import os
import re
import argparse

PATH = os.path.expanduser('~/Documents/')

def Title():
    print("==================================================================================")
    print("Automated Find and Replace Script")
    print("")
    print("Created By: Ashton Sidhu")
    print("Script that creates multiple files by replacing select words in a file template.")
    print("==================================================================================")
    print("")

def ReadRepFile(repFile):
    """
    Reads in the words to replace for each new file to be created

    String -> String

    Input: Words to be replace file name
    Ouput: List of words to replace for each file
    """
    try:
        with open(PATH + repFile + ".txt") as rFile:
            lines = rFile.readlines()
        lines = [x.strip().replace("=", ":").split(",") for x in lines]
        return lines

    except:
        print("Cannot parse word replacement file.")
        print("Please ensure it is following the file specifications described on the FindReplaceScript Github page README.")


def ReadTempFile(tempFile):
    """
    Reads in the content of the document of which the words are going to be replaced

    String -> String

    Input: Document file name
    Ouput: String content of file
   """
    try:
        doc = Document(PATH + tempFile + ".docx")

        docText = ""
        
        for para in doc.paragraphs:
            docText += (para.text + '\n')

        return docText

    except:
        print("Cannot read word file.")

def ListToDic(li):
    """
    Converts list of strings separated by '=' or ':' to a dictionary

    List[String] -> Dictionary

    Input: List of strings
    Ouput: Ditionary of words and their replacement
   """      
    
    replaceDict = [dict([x.split(':') for x in item]) for item in li]

    return replaceDict     

def TextReplace(doc, repDict):
    """
    Replaces the specified word with the specified replacement word.

    (String, Dictionary) -> String

    Input: String of document text, Dictionary of words to be replaced
    Ouput: String of document content with replaced words.
   """

    for key in repDict:
        
        doc = re.sub( "\\b" + key + "\\b" ,  repDict[key], doc )

    return doc

def WriteToFile(fileContent, index):
    """
    Writes string form of document to a .docx file.

    (String, Int) -> File

    Input: Document content, File processing number
    Ouput: .docx file.
   """

    doc = Document()
    doc.add_paragraph(fileContent)
    doc.save(PATH + 'file' + str(index) + '.docx')

def FileExists(fileName, format):
    """
    Check if file exists in My Documents folder.

    (String, String) -> Boolean

    Input: Filename, File Format
    Output: True if file exists, false otherwiese
    """

    if os.path.isfile(PATH + fileName + format):
        return True

    return False

def main(args):
    
    Title()
    #Collect initial parameters: template file name, word replace file
    print("Ensure your template file and replace file are in your Documents folder.")

    if (not args.doc and not args.temp):
        while True:
            tempFileName = input("Enter template file: ")

            if FileExists(tempFileName, ".docx"):
                break
            else:
                print(" File does not exist.")

        while True:
            repFileName = input("Enter number of words to replace file: ")

            if FileExists(repFileName, ".txt"):
                break
            else:
                print(" File does not exist.")
    else:
        repFileName = args.temp[0]
        tempFileName = args.doc[0]

    #Read and parse words to replace and how many files to create.
    repContent = ReadRepFile(repFileName)
    docContent = ReadTempFile(tempFileName)

    #Convert list of words to replace to a dictionary
    repContent = ListToDic(repContent)
   
    #Replace words in the document with the words specified from the user
    print()
    for idx, dic in enumerate(repContent):
        fileContent = TextReplace(docContent, dic)
        WriteToFile(fileContent, idx)
        print("File " + str(idx) + " has been created.")

    print("")
    print("All files have been created successfully.")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Find and Replace words and phrases in a template document')
    parser.add_argument('--doc', metavar='file', nargs='+', help='Template Document')
    parser.add_argument('--temp', metavar='file', help='File of words to replace for each document')
    args = parser.parse_args()

    main(args)
    